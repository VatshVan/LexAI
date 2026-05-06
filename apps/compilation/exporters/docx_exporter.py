from pathlib import Path
from uuid import UUID
from django.conf import settings
from docx import Document
from docx.shared import RGBColor, Pt
from .base import BaseExporter
from ..models import CompiledDocument, ReviewChecklistItem


class DOCXExporter(BaseExporter):
    def export(self, document_id: UUID) -> Path:
        doc_model = CompiledDocument.objects.get(document_id=document_id)
        self._check_permission(doc_model)

        if doc_model.export_docx_path:
            p = Path(doc_model.export_docx_path)
            if p.exists():
                return p

        out_dir = Path(settings.EXPORT_ROOT) / str(document_id)
        out_dir.mkdir(parents=True, exist_ok=True)
        out_path = out_dir / f"{document_id}.docx"

        doc = Document()
        doc.add_heading("LexAI — AI-Assisted Legal Document", 0)
        doc.add_paragraph(f"Template: {doc_model.template_name or 'Analysis'}")
        doc.add_paragraph(f"Verification Score: {doc_model.review_completion_pct:.0f}%")
        doc.add_paragraph("")

        for item in ReviewChecklistItem.objects.filter(document=doc_model).order_by("clause_index"):
            p = doc.add_paragraph()
            run = p.add_run(item.clause_text)
            if item.is_null_field:
                run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
                run.bold = True
                p.add_run(" [REQUIRED — Fill before filing]").font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
            elif item.verification_verdict == "INSUFFICIENT":
                run.font.highlight_color = 7  # Yellow
                doc.add_paragraph(f"⚠ {item.citation_string or 'Verify manually'}").runs[0].font.size = Pt(9)
            elif item.citation_string:
                doc.add_paragraph(item.citation_string).runs[0].font.size = Pt(9)

        doc.add_paragraph("")
        disc = doc.add_paragraph(
            "AI-assisted document. Review all flagged content before filing.")
        disc.runs[0].font.size = Pt(9)
        disc.runs[0].italic = True

        doc.save(str(out_path))
        doc_model.export_docx_path = str(out_path)
        doc_model.save(update_fields=["export_docx_path"])
        return out_path
