import os

files = {
    'apps/verification/__init__.py': '',
    
    'apps/verification/schemas.py': '''from pydantic import BaseModel, Field
from enum import Enum
from uuid import UUID, uuid4

class VerificationVerdict(str, Enum):
    VERIFIED             = "VERIFIED"
    INSUFFICIENT         = "INSUFFICIENT"
    HALLUCINATED         = "HALLUCINATED"
    EXTERNALLY_VERIFIED  = "EXTERNALLY_VERIFIED"
    EXTERNALLY_FAILED    = "EXTERNALLY_FAILED"
    PURGED               = "PURGED"

class AtomicClaim(BaseModel):
    atomic_claim_id: UUID = Field(default_factory=uuid4)
    parent_claim_id: UUID
    atomic_text: str
    subject: str
    predicate: str
    requires_external_check: bool = False

class EntailmentResult(BaseModel):
    atomic_claim_id: UUID
    source_vector_id: str
    source_text: str
    entailment_score: float
    contradiction_score: float
    neutral_score: float
    verdict: VerificationVerdict
    scorer_method: str

class WebResearchResult(BaseModel):
    atomic_claim_id: UUID
    query_used: str
    source_url: str
    retrieved_text: str
    supports_claim: bool
    verdict: VerificationVerdict

class ClaimVerificationSummary(BaseModel):
    claim_id: UUID
    atomic_claims: list[AtomicClaim]
    entailment_results: list[EntailmentResult]
    web_research_results: list[WebResearchResult]
    final_verdict: VerificationVerdict
    final_score: float
    citation_string: str
    purge_reason: str | None

class VerificationReport(BaseModel):
    query_id: UUID
    total_claims_received: int
    total_claims_verified: int
    total_claims_purged: int
    total_claims_externally_verified: int
    verification_score: float
    claim_summaries: list[ClaimVerificationSummary]
    verified_claims: list[UUID]
    purged_claims: list[UUID]
    verification_duration_ms: int
    agents_used: list[str]
''',

    'apps/claude_client/prompts/claim_decomposition.txt': '''Break a compound legal statement into atomic verifiable units.
Rules: 1 fact per atomic claim. Replace pronouns with named entities.
Set requires_external_check=true if claim cites a statute or case name.
Max 3 atomics per input claim. If already atomic, return as single-item list.
Respond ONLY with valid JSON:
{"atomic_claims":[{"atomic_text":"...","subject":"...","predicate":"...","requires_external_check":false}]}
''',

    'apps/claude_client/prompts/entailment_scoring.txt': '''Determine if source text ENTAILS, CONTRADICTS, or is NEUTRAL to the claim.
ENTAILS: claim follows necessarily from source (exact names/dates/numbers must match).
CONTRADICTS: source makes claim false.
NEUTRAL: related but not sufficient to confirm claim.
Scores must sum to 1.0. Be strict — paraphrase != entailment.
Respond ONLY with valid JSON:
{"entailment_score":0.0,"contradiction_score":0.0,"neutral_score":0.0,"reasoning":"one sentence"}
''',

    'apps/claude_client/prompts/web_research_verdict.txt': '''Does the retrieved legal text confirm the claim?
Only confirm if retrieved text explicitly addresses same statute/section/point.
Same statute, different subsection = NOT confirmed.
Respond ONLY with valid JSON:
{"supports_claim":false,"confidence":0.0,"reasoning":"one sentence"}
''',

    'apps/verification/agents/claim_parser.py': '''"""
Cost: Haiku, ~200 input + 100 output per batch of 10 claims ≈ $0.0003
Fast path: template_extraction claims skip Claude (already atomic).
"""
from uuid import UUID, uuid4
from pydantic import BaseModel
from apps.agents.base import BaseAgent
from apps.agents.schemas import FactualClaim
from apps.verification.schemas import AtomicClaim
from apps.claude_client.client import get_claude
import json

class _AtomicResponse(BaseModel):
    atomic_claims: list[dict]


class ClaimParserAgent(BaseAgent):
    agent_name = "ClaimParser"

    def _execute(self, claims: list[FactualClaim]) -> list[AtomicClaim]:
        result = []
        batch, batch_ids = [], []

        for claim in claims:
            if claim.claim_type == "template_extraction":
                # Fast path: already atomic
                result.append(AtomicClaim(
                    parent_claim_id=claim.claim_id,
                    atomic_text=claim.claim_text,
                    subject=claim.claim_text.split(":")[0] if ":" in claim.claim_text else "",
                    predicate=claim.claim_text.split(":", 1)[1].strip() if ":" in claim.claim_text else claim.claim_text,
                    requires_external_check=False,
                ))
            else:
                batch.append(f"CLAIM_ID:{claim.claim_id}\\n{claim.claim_text}")
                batch_ids.append(claim.claim_id)

            if len(batch) == 10:
                result.extend(self._call_claude(batch, batch_ids))
                batch, batch_ids = [], []

        if batch:
            result.extend(self._call_claude(batch, batch_ids))
        return result

    def _call_claude(self, batch: list[str], ids: list[UUID]) -> list[AtomicClaim]:
        user_msg = "DECOMPOSE EACH:\\n\\n" + "\\n---\\n".join(batch)
        try:
            raw = get_claude().haiku("claim_decomposition", user_msg)
            clean = raw.strip().removeprefix("```json").removeprefix("```").removesuffix("```").strip()
            data = json.loads(clean)
            atomics = data.get("atomic_claims", [])
        except Exception:
            # Fallback: treat each claim as its own atomic
            return [AtomicClaim(parent_claim_id=cid, atomic_text=b.split("\\n", 1)[1],
                                subject="", predicate="", requires_external_check=False)
                    for b, cid in zip(batch, ids)]

        result = []
        for i, a in enumerate(atomics):
            parent_id = ids[i] if i < len(ids) else ids[-1]
            result.append(AtomicClaim(
                parent_claim_id=parent_id,
                atomic_text=a.get("atomic_text", ""),
                subject=a.get("subject", ""),
                predicate=a.get("predicate", ""),
                requires_external_check=a.get("requires_external_check", False),
            ))
        return result
''',

    'apps/verification/agents/entailment_checker.py': '''"""
Cost optimization:
  Cosine fast-path (no API call): handles ~70% of claims
  Sonnet slow-path: only for 0.45 < cosine < 0.82
  Typical cost: ~1-2 Sonnet calls per query ≈ $0.003
"""
import numpy as np
from uuid import UUID
from pydantic import BaseModel
from apps.agents.base import BaseAgent
from apps.agents.schemas import FactualClaim
from apps.verification.schemas import AtomicClaim, EntailmentResult, VerificationVerdict
from apps.documents.services.embedding import EmbeddingService
from apps.vector_store.repository import VectorRepository
from apps.claude_client.client import get_claude
from django.conf import settings
import json


class _EntailmentResponse(BaseModel):
    entailment_score: float
    contradiction_score: float
    neutral_score: float
    reasoning: str


def _cosine(a, b):
    a, b = np.array(a), np.array(b)
    return float(np.dot(a, b) / (np.linalg.norm(a) * np.linalg.norm(b) + 1e-9))


class EntailmentCheckerAgent(BaseAgent):
    agent_name = "EntailmentChecker"

    def _execute(
        self,
        atomic_claims: list[AtomicClaim],
        original_claims: list[FactualClaim],
    ) -> list[EntailmentResult]:
        claim_map = {c.claim_id: c for c in original_claims}
        embed = EmbeddingService()
        repo = VectorRepository()
        results = []

        HI = settings.COSINE_VERIFIED_THRESHOLD
        LO = settings.COSINE_HALLUCINATED_THRESHOLD

        for ac in atomic_claims:
            parent = claim_map.get(ac.parent_claim_id)
            if not parent:
                continue
            query_vec = embed.embed_query(ac.atomic_text)

            for vid in parent.source_vector_ids:
                vdoc = repo.get_by_vector_id(str(parent.source_vector_ids[0].split(":")[0]), vid)
                if not vdoc:
                    continue
                # Cosine fast-path
                if vdoc.embedding:
                    cosine = _cosine(query_vec, vdoc.embedding)
                    if cosine >= HI:
                        verdict = VerificationVerdict.VERIFIED
                        e, c, n = cosine, 0.0, 1-cosine
                        method = "cosine"
                    elif cosine <= LO:
                        verdict = VerificationVerdict.HALLUCINATED
                        e, c, n = 0.0, cosine, 1-cosine
                        method = "cosine"
                    else:
                        # Slow-path: Claude entailment
                        user_msg = f"CLAIM: {ac.atomic_text}\\nSOURCE: {vdoc.text}"
                        try:
                            resp = get_claude().sonnet_json(
                                "entailment_scoring", user_msg, _EntailmentResponse)
                            e, c, n = resp.entailment_score, resp.contradiction_score, resp.neutral_score
                        except Exception:
                            e, c, n = cosine, 0.0, 1-cosine
                        verdict = (VerificationVerdict.VERIFIED if e >= 0.70
                                   else VerificationVerdict.HALLUCINATED if e <= 0.30
                                   else VerificationVerdict.INSUFFICIENT)
                        method = "cross_encoder"
                else:
                    e, c, n = 0.5, 0.0, 0.5
                    verdict = VerificationVerdict.INSUFFICIENT
                    method = "no_embedding"

                results.append(EntailmentResult(
                    atomic_claim_id=ac.atomic_claim_id,
                    source_vector_id=vid,
                    source_text=vdoc.text[:300],
                    entailment_score=e,
                    contradiction_score=c,
                    neutral_score=n,
                    verdict=verdict,
                    scorer_method=method,
                ))
        return results
''',

    'apps/verification/agents/web_research_agent.py': '''"""
Cost: Haiku for verdict ≈ $0.0001 per claim checked
Only runs on claims that are INSUFFICIENT AND require_external_check=True
Typical: 0-2 claims per query ≈ $0.0002
"""
import httpx
import re
import json
from pydantic import BaseModel
from apps.agents.base import BaseAgent
from apps.verification.schemas import AtomicClaim, EntailmentResult, WebResearchResult, VerificationVerdict
from apps.claude_client.client import get_claude

WHITELIST = {
    "indiankanoon.org", "legislative.gov.in", "main.sci.gov.in",
    "districts.ecourts.gov.in", "bombayhighcourt.nic.in",
    "delhihighcourt.nic.in", "barandbench.com", "livelaw.in",
}

class _VerdictResponse(BaseModel):
    supports_claim: bool
    confidence: float
    reasoning: str


class WebResearchAgent(BaseAgent):
    agent_name = "WebResearch"

    def _execute(
        self,
        atomic_claims: list[AtomicClaim],
        entailment_results: list[EntailmentResult],
    ) -> list[WebResearchResult]:
        ent_map = {r.atomic_claim_id: r for r in entailment_results}
        results = []

        eligible = [
            ac for ac in atomic_claims
            if ac.requires_external_check
            and ent_map.get(ac.atomic_claim_id, None) is not None
            and ent_map[ac.atomic_claim_id].verdict == VerificationVerdict.INSUFFICIENT
        ]

        for ac in eligible:
            result = self._research_claim(ac)
            if result:
                results.append(result)
        return results

    def _research_claim(self, ac: AtomicClaim) -> WebResearchResult | None:
        import urllib.parse
        query = urllib.parse.quote(ac.atomic_text[:100])
        url = f"https://indiankanoon.org/search/?formInput={query}"

        try:
            with httpx.Client(timeout=10.0, follow_redirects=True,
                              headers={"User-Agent": "LexAI-Research/1.0"}) as client:
                resp = client.get(url)
            text = re.sub(r"<[^>]+>", " ", resp.text)
            text = re.sub(r"\\s+", " ", text).strip()[:500]
            if len(text) < 50:
                return None
        except Exception:
            return None

        user_msg = f"CLAIM: {ac.atomic_text}\\nRETRIEVED: {text}\\nURL: {url}"
        try:
            verdict_resp = get_claude().haiku_json("web_research_verdict", user_msg, _VerdictResponse)
        except Exception:
            return None

        return WebResearchResult(
            atomic_claim_id=ac.atomic_claim_id,
            query_used=ac.atomic_text[:100],
            source_url=url,
            retrieved_text=text,
            supports_claim=verdict_resp.supports_claim,
            verdict=VerificationVerdict.EXTERNALLY_VERIFIED
            if verdict_resp.supports_claim else VerificationVerdict.EXTERNALLY_FAILED,
        )
''',

    'apps/verification/agents/resolution_gate.py': '''import statistics
from uuid import UUID
from apps.agents.base import BaseAgent
from apps.agents.schemas import FactualClaim
from apps.verification.schemas import (
    AtomicClaim, EntailmentResult, WebResearchResult,
    ClaimVerificationSummary, VerificationVerdict,
)
from django.conf import settings

E_WEIGHT = 0.70
W_WEIGHT = 0.30


class ResolutionGate(BaseAgent):
    agent_name = "ResolutionGate"

    def _execute(
        self,
        original_claims: list[FactualClaim],
        atomic_claims: list[AtomicClaim],
        entailment_results: list[EntailmentResult],
        web_research_results: list[WebResearchResult],
    ) -> tuple[list[FactualClaim], list[ClaimVerificationSummary]]:
        ent_by_ac = {}
        for e in entailment_results:
            ent_by_ac.setdefault(e.atomic_claim_id, []).append(e)
        web_by_ac = {w.atomic_claim_id: w for w in web_research_results}
        ac_by_parent = {}
        for ac in atomic_claims:
            ac_by_parent.setdefault(ac.parent_claim_id, []).append(ac)

        verified, summaries = [], []
        PURGE = settings.PURGE_SCORE_THRESHOLD
        CONTRA = settings.CONTRADICTION_PURGE_THRESHOLD

        for claim in original_claims:
            atomics = ac_by_parent.get(claim.claim_id, [])
            if not atomics:
                # No atomics generated — keep as insufficient
                summaries.append(ClaimVerificationSummary(
                    claim_id=claim.claim_id,
                    atomic_claims=[], entailment_results=[], web_research_results=[],
                    final_verdict=VerificationVerdict.INSUFFICIENT,
                    final_score=0.5, citation_string="", purge_reason=None,
                ))
                claim.is_flagged = True
                verified.append(claim)
                continue

            atomic_scores = []
            direct_contradiction = False
            best_source = None

            for ac in atomics:
                ents = ent_by_ac.get(ac.atomic_claim_id, [])
                if not ents:
                    atomic_scores.append(0.5)
                    continue
                best_ent = max(ents, key=lambda e: e.entailment_score)
                if best_ent.contradiction_score > CONTRA:
                    direct_contradiction = True
                    break
                if not best_source or best_ent.entailment_score > (best_source.entailment_score if best_source else 0):
                    best_source = best_ent

                web = web_by_ac.get(ac.atomic_claim_id)
                if web:
                    score = best_ent.entailment_score * E_WEIGHT + (1.0 if web.supports_claim else 0.0) * W_WEIGHT
                else:
                    score = best_ent.entailment_score
                atomic_scores.append(score)

            if direct_contradiction:
                verdict = VerificationVerdict.PURGED
                final_score = 0.0
                purge_reason = "direct_contradiction"
            else:
                final_score = statistics.mean(atomic_scores) if atomic_scores else 0.0
                # NULL template fields: never purge
                if claim.claim_type == "template_extraction" and claim.claim_text.endswith("NULL"):
                    verdict = VerificationVerdict.INSUFFICIENT
                    purge_reason = None
                elif final_score >= 0.70:
                    verdict = VerificationVerdict.VERIFIED
                    purge_reason = None
                elif final_score >= 0.50:
                    verdict = VerificationVerdict.INSUFFICIENT
                    purge_reason = None
                else:
                    verdict = VerificationVerdict.PURGED
                    purge_reason = f"low_composite_score:{final_score:.2f}"

            # Citation string from best source
            citation = ""
            if best_source and verdict != VerificationVerdict.PURGED:
                m = best_source.source_text
                citation = f"[Src: {best_source.source_vector_id[:8]}]"

            summaries.append(ClaimVerificationSummary(
                claim_id=claim.claim_id,
                atomic_claims=atomics,
                entailment_results=ent_by_ac.get(
                    atomics[0].atomic_claim_id, []) if atomics else [],
                web_research_results=[web_by_ac[ac.atomic_claim_id]
                                       for ac in atomics if ac.atomic_claim_id in web_by_ac],
                final_verdict=verdict,
                final_score=final_score,
                citation_string=citation,
                purge_reason=purge_reason,
            ))

            if verdict != VerificationVerdict.PURGED:
                claim.is_verified = (verdict == VerificationVerdict.VERIFIED)
                claim.is_flagged = (verdict == VerificationVerdict.INSUFFICIENT)
                claim.verification_score = final_score
                verified.append(claim)

        return verified, summaries
''',

    'apps/verification/engine.py': '''from uuid import UUID
from time import perf_counter
from apps.agents.schemas import FactualClaim
from apps.verification.agents.claim_parser import ClaimParserAgent
from apps.verification.agents.entailment_checker import EntailmentCheckerAgent
from apps.verification.agents.web_research_agent import WebResearchAgent
from apps.verification.agents.resolution_gate import ResolutionGate
from apps.verification.schemas import VerificationReport, VerificationVerdict
from apps.orchestration.context_manager import SSEPublisher
import structlog

log = structlog.get_logger()


class VerificationEngine:
    def run(
        self,
        query_id: UUID,
        session_id: UUID,
        claims: list[FactualClaim],
    ) -> tuple[list[FactualClaim], VerificationReport]:
        sse = SSEPublisher()
        start = perf_counter()
        qid = str(query_id)
        sse.publish(qid, {"event": "verification_start", "total_claims": len(claims)})

        atomic_claims = ClaimParserAgent().execute(claims)
        sse.publish(qid, {"event": "agent_complete", "agent": "ClaimParser",
                          "atomic_claims": len(atomic_claims)})

        entailment_results = EntailmentCheckerAgent().execute(
            atomic_claims=atomic_claims, original_claims=claims)
        sse.publish(qid, {"event": "agent_complete", "agent": "EntailmentChecker"})

        eligible_for_web = [
            ac for ac in atomic_claims
            if ac.requires_external_check
            and any(e.atomic_claim_id == ac.atomic_claim_id
                    and e.verdict.value in ("INSUFFICIENT",)
                    for e in entailment_results)
        ]
        agents_used = ["ClaimParser", "EntailmentChecker"]
        if eligible_for_web:
            web_results = WebResearchAgent().execute(
                atomic_claims=eligible_for_web,
                entailment_results=entailment_results,
            )
            agents_used.append("WebResearch")
            sse.publish(qid, {"event": "agent_complete", "agent": "WebResearch",
                              "checked": len(web_results)})
        else:
            web_results = []
            sse.publish(qid, {"event": "agent_skipped", "agent": "WebResearch"})

        verified_claims, summaries = ResolutionGate().execute(
            original_claims=claims,
            atomic_claims=atomic_claims,
            entailment_results=entailment_results,
            web_research_results=web_results,
        )
        agents_used.append("ResolutionGate")

        purged = [s for s in summaries if s.final_verdict == VerificationVerdict.PURGED]
        ev_count = sum(1 for s in summaries
                       if s.final_verdict == VerificationVerdict.EXTERNALLY_VERIFIED)

        report = VerificationReport(
            query_id=query_id,
            total_claims_received=len(claims),
            total_claims_verified=len([s for s in summaries
                                        if s.final_verdict == VerificationVerdict.VERIFIED]),
            total_claims_purged=len(purged),
            total_claims_externally_verified=ev_count,
            verification_score=len(verified_claims) / max(len(claims), 1),
            claim_summaries=summaries,
            verified_claims=[c.claim_id for c in verified_claims],
            purged_claims=[s.claim_id for s in purged],
            verification_duration_ms=int((perf_counter() - start) * 1000),
            agents_used=agents_used,
        )
        sse.publish(qid, {"event": "verification_complete",
                          "score": report.verification_score,
                          "passed": len(verified_claims),
                          "purged": len(purged)})
        return verified_claims, report
''',

    'apps/compilation/assembler.py': '''from uuid import UUID
from apps.agents.schemas import FactualClaim, DraftingOutput, SynthesisOutput
from apps.verification.schemas import ClaimVerificationSummary, VerificationVerdict
from .models import CompiledDocument, ReviewChecklistItem
from apps.orchestration.models import QuerySession


class AssemblerService:
    def assemble(
        self,
        query_id: UUID,
        verified_claims: list[FactualClaim],
        summaries: list[ClaimVerificationSummary],
        result_payload,
    ) -> CompiledDocument:
        qs = QuerySession.objects.get(query_id=query_id)
        summary_map = {str(s.claim_id): s for s in summaries}

        is_drafting = isinstance(result_payload, DraftingOutput)

        doc = CompiledDocument.objects.create(
            query=qs,
            status=CompiledDocument.Status.ASSEMBLING,
            template_name=result_payload.template_name if is_drafting else "",
            assembled_html="",
            total_clauses=len(verified_claims),
        )

        items, html_parts = [], []
        for idx, claim in enumerate(verified_claims):
            s = summary_map.get(str(claim.claim_id))
            verdict = s.final_verdict.value if s else "UNKNOWN"
            citation = s.citation_string if s else ""
            score = s.final_score if s else 0.0

            is_null = is_drafting and claim.claim_text.split(":", 1)[-1].strip() == "None"
            css_class = "null-field" if is_null else (
                "verified" if verdict == "VERIFIED" else "insufficient")

            html_parts.append(
                f'<div class="clause {css_class}" data-claim-id="{claim.claim_id}" '
                f'data-vector-ids="{",".join(claim.source_vector_ids)}" '
                f'data-score="{score:.2f}">'
                f'<p class="clause-text">{claim.claim_text}</p>'
                f'{"<span class=\\"citation\\">" + citation + "</span>" if citation else ""}'
                f'{"<p class=\\"flag\\">⚠ Verify manually — evidence inconclusive</p>" if verdict == "INSUFFICIENT" else ""}'
                f'{"<p class=\\"null-notice\\">[REQUIRED — Fill before filing]</p>" if is_null else ""}'
                f'</div>'
            )
            items.append(ReviewChecklistItem(
                document=doc, clause_index=idx,
                clause_text=claim.claim_text,
                citation_string=citation,
                source_vector_ids=claim.source_vector_ids,
                verification_verdict=verdict,
                verification_score=score,
                is_null_field=is_null,
                is_approved=False,
            ))

        ReviewChecklistItem.objects.bulk_create(items)
        doc.assembled_html = (
            '<div class="lexai-analysis-result">' + "".join(html_parts) + "</div>"
        )
        doc.status = CompiledDocument.Status.PENDING_REVIEW
        doc.total_clauses = len(items)
        doc.save()
        return doc
''',

    'apps/verification/signals.py': '''from django.dispatch import receiver
from apps.orchestration.pipeline import pre_finalize_signal
from apps.verification.engine import VerificationEngine
from apps.compilation.assembler import AssemblerService


@receiver(pre_finalize_signal)
def run_verification(sender, query_id, session_id, result_payload, query_session, **kwargs):
    claims = getattr(result_payload, "claims", None)
    if not claims:
        return
    verified_claims, report = VerificationEngine().run(
        query_id=query_id,
        session_id=session_id,
        claims=claims,
    )
    # Write report to QuerySession
    query_session.verification_report = report.model_dump(mode="json")
    query_session.save(update_fields=["verification_report"])

    # Replace claims in payload with verified subset
    result_payload.claims = verified_claims

    # Trigger compilation
    AssemblerService().assemble(
        query_id=query_id,
        verified_claims=verified_claims,
        summaries=report.claim_summaries,
        result_payload=result_payload,
    )
''',

    'apps/verification/apps.py': '''from django.apps import AppConfig

class VerificationConfig(AppConfig):
    name = "apps.verification"

    def ready(self):
        import apps.verification.signals  # noqa: connect receivers
''',

    'apps/compilation/models.py': '''from django.db import models
import uuid

class CompiledDocument(models.Model):
    class Status(models.TextChoices):
        ASSEMBLING      = "ASSEMBLING"
        PENDING_REVIEW  = "PENDING_REVIEW"
        REVIEW_COMPLETE = "REVIEW_COMPLETE"
        EXPORTED        = "EXPORTED"

    document_id           = models.UUIDField(primary_key=True, default=uuid.uuid4)
    query                 = models.OneToOneField(
        "orchestration.QuerySession", on_delete=models.CASCADE,
        related_name="compiled_document"
    )
    status                = models.CharField(max_length=20, default=Status.ASSEMBLING)
    template_name         = models.CharField(max_length=100, blank=True)
    assembled_html        = models.TextField()
    review_completion_pct = models.FloatField(default=0.0)
    total_clauses         = models.PositiveIntegerField(default=0)
    approved_clauses      = models.PositiveIntegerField(default=0)
    export_pdf_path       = models.CharField(max_length=500, blank=True)
    export_docx_path      = models.CharField(max_length=500, blank=True)
    exported_at           = models.DateTimeField(null=True)
    created_at            = models.DateTimeField(auto_now_add=True)


class ReviewChecklistItem(models.Model):
    document            = models.ForeignKey(CompiledDocument, on_delete=models.CASCADE,
                                            related_name="review_items")
    clause_index        = models.PositiveIntegerField()
    clause_text         = models.TextField()
    citation_string     = models.CharField(max_length=500, blank=True)
    source_vector_ids   = models.JSONField(default=list)
    verification_verdict= models.CharField(max_length=30, blank=True)
    verification_score  = models.FloatField(default=0.0)
    is_null_field       = models.BooleanField(default=False)
    is_approved         = models.BooleanField(default=False)
    approved_at         = models.DateTimeField(null=True)

    class Meta:
        unique_together = [("document", "clause_index")]
        ordering = ["clause_index"]
''',

    'apps/compilation/review_manager.py': '''from django.utils import timezone
from .models import CompiledDocument, ReviewChecklistItem


class ReviewManager:
    def approve_clause(self, document_id, clause_index: int) -> dict:
        doc = CompiledDocument.objects.get(document_id=document_id)
        item = ReviewChecklistItem.objects.get(document=doc, clause_index=clause_index)
        if item.is_null_field:
            raise ValueError("NULL fields cannot be auto-approved.")
        item.is_approved = True
        item.approved_at = timezone.now()
        item.save()
        return self._recompute(doc)

    def approve_all(self, document_id) -> dict:
        doc = CompiledDocument.objects.get(document_id=document_id)
        ReviewChecklistItem.objects.filter(
            document=doc, is_null_field=False, is_approved=False
        ).update(is_approved=True, approved_at=timezone.now())
        return self._recompute(doc)

    def reset(self, document_id) -> dict:
        doc = CompiledDocument.objects.get(document_id=document_id)
        ReviewChecklistItem.objects.filter(document=doc).update(is_approved=False, approved_at=None)
        return self._recompute(doc)

    def _recompute(self, doc: CompiledDocument) -> dict:
        total = doc.review_items.count()
        non_null = doc.review_items.filter(is_null_field=False)
        approved = non_null.filter(is_approved=True).count()
        null_count = doc.review_items.filter(is_null_field=True).count()
        pct = (approved / non_null.count() * 100) if non_null.count() > 0 else 0.0
        can_export = pct >= 100.0
        status = (CompiledDocument.Status.REVIEW_COMPLETE
                  if can_export else CompiledDocument.Status.PENDING_REVIEW)
        doc.approved_clauses = approved
        doc.review_completion_pct = pct
        doc.status = status
        doc.save(update_fields=["approved_clauses", "review_completion_pct", "status"])
        return {"approved_clauses": approved, "total_clauses": total,
                "null_clauses": null_count, "review_completion_pct": pct,
                "can_export": can_export, "status": status}
''',

    'apps/compilation/exporters/base.py': '''from abc import ABC, abstractmethod
from pathlib import Path
from uuid import UUID
from ..models import CompiledDocument


class ExportNotPermittedError(Exception): pass


class BaseExporter(ABC):
    def _check_permission(self, doc: CompiledDocument):
        if doc.status != CompiledDocument.Status.REVIEW_COMPLETE:
            pending = doc.total_clauses - doc.approved_clauses
            raise ExportNotPermittedError(
                f"Export blocked: {pending} clause(s) pending approval."
            )

    @abstractmethod
    def export(self, document_id: UUID) -> Path: ...
''',

    'apps/compilation/exporters/pdf_exporter.py': '''from pathlib import Path
from uuid import UUID
from django.conf import settings
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from .base import BaseExporter
from ..models import CompiledDocument, ReviewChecklistItem


class PDFExporter(BaseExporter):
    def export(self, document_id: UUID) -> Path:
        doc = CompiledDocument.objects.get(document_id=document_id)
        self._check_permission(doc)

        if doc.export_pdf_path:
            p = Path(doc.export_pdf_path)
            if p.exists():
                return p

        out_dir = Path(settings.EXPORT_ROOT) / str(document_id)
        out_dir.mkdir(parents=True, exist_ok=True)
        out_path = out_dir / f"{document_id}.pdf"

        styles = getSampleStyleSheet()
        citation_style = ParagraphStyle("citation", parent=styles["Normal"],
                                        textColor=colors.grey, fontSize=8)
        flag_style = ParagraphStyle("flag", parent=styles["Normal"],
                                    textColor=colors.orange)
        null_style = ParagraphStyle("null", parent=styles["Normal"],
                                    textColor=colors.red, fontName="Helvetica-Bold")

        flowables = [
            Paragraph("LexAI — AI-Assisted Legal Document", styles["Title"]),
            Paragraph(f"Template: {doc.template_name or 'Analysis'}", styles["Normal"]),
            Paragraph(f"Verification Score: {doc.review_completion_pct:.0f}%", styles["Normal"]),
            Spacer(1, 20),
        ]

        for item in ReviewChecklistItem.objects.filter(document=doc).order_by("clause_index"):
            if item.is_null_field:
                flowables.append(Paragraph(
                    f"{item.clause_text} [REQUIRED — FILL BEFORE FILING]", null_style))
            else:
                flowables.append(Paragraph(item.clause_text, styles["Normal"]))
                if item.citation_string:
                    flowables.append(Paragraph(item.citation_string, citation_style))
                if item.verification_verdict == "INSUFFICIENT":
                    flowables.append(Paragraph("⚠ Verify manually", flag_style))
            flowables.append(Spacer(1, 8))

        flowables.append(Spacer(1, 30))
        flowables.append(Paragraph(
            "AI-assisted document. Review all flagged content before filing.",
            ParagraphStyle("disclaimer", parent=styles["Normal"],
                           textColor=colors.grey, fontSize=8, italics=1)
        ))

        pdf = SimpleDocTemplate(str(out_path), pagesize=A4)
        pdf.build(flowables)
        doc.export_pdf_path = str(out_path)
        doc.save(update_fields=["export_pdf_path"])
        return out_path
''',

    'apps/compilation/exporters/docx_exporter.py': '''from pathlib import Path
from uuid import UUID
from django.conf import settings
from docx import Document
from docx.shared import RGBColor, Pt
from .base import BaseExporter
from ..models import CompiledDocument, ReviewChecklistItem


class DOCXExporter(BaseExporter):
    def export(self, document_id: UUID) -> Path:
        doc_model = CompiledDocument.objects.get(document_id=document_id)
        self._check_permission(doc_model)

        if doc_model.export_docx_path:
            p = Path(doc_model.export_docx_path)
            if p.exists():
                return p

        out_dir = Path(settings.EXPORT_ROOT) / str(document_id)
        out_dir.mkdir(parents=True, exist_ok=True)
        out_path = out_dir / f"{document_id}.docx"

        doc = Document()
        doc.add_heading("LexAI — AI-Assisted Legal Document", 0)
        doc.add_paragraph(f"Template: {doc_model.template_name or 'Analysis'}")
        doc.add_paragraph(f"Verification Score: {doc_model.review_completion_pct:.0f}%")
        doc.add_paragraph("")

        for item in ReviewChecklistItem.objects.filter(document=doc_model).order_by("clause_index"):
            p = doc.add_paragraph()
            run = p.add_run(item.clause_text)
            if item.is_null_field:
                run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
                run.bold = True
                p.add_run(" [REQUIRED — Fill before filing]").font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
            elif item.verification_verdict == "INSUFFICIENT":
                run.font.highlight_color = 7  # Yellow
                doc.add_paragraph(f"⚠ {item.citation_string or 'Verify manually'}").runs[0].font.size = Pt(9)
            elif item.citation_string:
                doc.add_paragraph(item.citation_string).runs[0].font.size = Pt(9)

        doc.add_paragraph("")
        disc = doc.add_paragraph(
            "AI-assisted document. Review all flagged content before filing.")
        disc.runs[0].font.size = Pt(9)
        disc.runs[0].italic = True

        doc.save(str(out_path))
        doc_model.export_docx_path = str(out_path)
        doc_model.save(update_fields=["export_docx_path"])
        return out_path
''',

    'apps/compilation/views.py': '''from django.http import FileResponse
from rest_framework.views import APIView
from rest_framework.response import Response
from .models import CompiledDocument, ReviewChecklistItem
from .review_manager import ReviewManager
from .exporters.base import ExportNotPermittedError
from .exporters.pdf_exporter import PDFExporter
from .exporters.docx_exporter import DOCXExporter
from pathlib import Path


def _checklist_data(doc):
    items = list(doc.review_items.values(
        "clause_index", "clause_text", "citation_string", "source_vector_ids",
        "verification_verdict", "verification_score", "is_null_field",
        "is_approved", "approved_at",
    ))
    return {
        "document_id": str(doc.document_id),
        "query_id": str(doc.query_id),
        "status": doc.status,
        "template_name": doc.template_name,
        "assembled_html": doc.assembled_html,
        "review_completion_pct": doc.review_completion_pct,
        "total_clauses": doc.total_clauses,
        "approved_clauses": doc.approved_clauses,
        "can_export": doc.status == CompiledDocument.Status.REVIEW_COMPLETE,
        "checklist": items,
    }


class CompiledDocumentView(APIView):
    def get(self, request, query_id):
        try:
            doc = CompiledDocument.objects.get(query__query_id=query_id)
        except CompiledDocument.DoesNotExist:
            return Response({"success": False, "error": "Not found"}, status=404)
        return Response({"success": True, "data": _checklist_data(doc)})


class ApproveClauseView(APIView):
    def post(self, request, query_id, clause_index):
        doc = CompiledDocument.objects.get(query__query_id=query_id)
        try:
            data = ReviewManager().approve_clause(doc.document_id, clause_index)
            return Response({"success": True, "data": data})
        except ValueError as e:
            return Response({"success": False, "error": str(e)}, status=400)


class ApproveAllView(APIView):
    def post(self, request, query_id):
        doc = CompiledDocument.objects.get(query__query_id=query_id)
        data = ReviewManager().approve_all(doc.document_id)
        return Response({"success": True, "data": data})


class ResetReviewView(APIView):
    def post(self, request, query_id):
        doc = CompiledDocument.objects.get(query__query_id=query_id)
        data = ReviewManager().reset(doc.document_id)
        return Response({"success": True, "data": data})


class ExportStatusView(APIView):
    def get(self, request, query_id):
        doc = CompiledDocument.objects.get(query__query_id=query_id)
        items = doc.review_items
        return Response({"success": True, "data": {
            "can_export": doc.status == CompiledDocument.Status.REVIEW_COMPLETE,
            "review_completion_pct": doc.review_completion_pct,
            "blockers": {
                "pending_approvals": items.filter(is_null_field=False, is_approved=False).count(),
                "null_fields_unfilled": items.filter(is_null_field=True).count(),
            },
            "exports_available": {
                "pdf": {"available": bool(doc.export_pdf_path),
                        "generated_at": None},
                "docx": {"available": bool(doc.export_docx_path),
                         "generated_at": None},
            }
        }})


class ExportView(APIView):
    EXPORTERS = {"pdf": PDFExporter, "docx": DOCXExporter}
    MIME = {
        "pdf": "application/pdf",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }

    def get(self, request, query_id, fmt):
        if fmt not in self.EXPORTERS:
            return Response({"success": False, "error": "Unsupported format"}, status=400)
        try:
            doc = CompiledDocument.objects.get(query__query_id=query_id)
            path = self.EXPORTERS[fmt]().export(doc.document_id)
            return FileResponse(open(path, "rb"), content_type=self.MIME[fmt],
                                as_attachment=True, filename=path.name)
        except ExportNotPermittedError as e:
            pending = doc.review_items.filter(is_null_field=False, is_approved=False).count()
            return Response({"success": False, "error": str(e),
                             "data": {"pending_approvals": pending}}, status=403)
''',

    'apps/compilation/urls.py': '''from django.urls import path
from . import views

urlpatterns = [
    path("queries/<str:query_id>/document/", views.CompiledDocumentView.as_view()),
    path("queries/<str:query_id>/document/clauses/<int:clause_index>/approve/",
         views.ApproveClauseView.as_view()),
    path("queries/<str:query_id>/document/clauses/approve-all/",
         views.ApproveAllView.as_view()),
    path("queries/<str:query_id>/document/clauses/reset/",
         views.ResetReviewView.as_view()),
    path("queries/<str:query_id>/document/export/status/",
         views.ExportStatusView.as_view()),
    path("queries/<str:query_id>/document/export/<str:fmt>/",
         views.ExportView.as_view()),
]
'''
}

for k, v in files.items():
    os.makedirs(os.path.dirname(k) or '.', exist_ok=True)
    with open(k, 'w', encoding='utf-8') as f:
        f.write(v)

print("done")
